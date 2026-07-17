import os
import json
import uuid
import pymupdf
from docx import Document
from sqlalchemy.orm import Session
from app.models.resume import Resume
from app.core.exceptions import NotFoundException, BadRequestException, LLMAPIException, handle_llm_exception
from app.core.config import settings
from app.services.llm_service import LLMService
from loguru import logger


FONT_MAP = {
    "arial": "helv",
    "helvetica": "helv",
    "calibri": "helv",
    "segoe ui": "helv",
    "open sans": "helv",
    "roboto": "helv",
    "lato": "helv",
    "montserrat": "helv",
    "times new roman": "tiro",
    "times": "tiro",
    "georgia": "tiro",
    "garamond": "tiro",
    "palatino": "tiro",
    "courier new": "cour",
    "courier": "cour",
    "consolas": "cour",
    "monaco": "cour",
}

AI_SYSTEM_PROMPT = """You are a professional resume editor. Your ONLY job is to improve the text quality of resume content.

RULES - You MUST follow these exactly:
1. ONLY correct: grammar, spelling, punctuation, sentence structure
2. ONLY improve: weak bullet points, action verbs, ATS keywords, readability, professional tone. DIVERSIFY action verbs: Do not repeat the same action verb (like "Developed" or "Built") across different bullet points. Use a variety of strong, professional Software Engineering action verbs such as: Engineered, Architected, Implemented, Spearheaded, Optimized, Formulated, Orchestrated, Designed, Devised, Automated, Modernized, Refactored.
3. NEVER change: the meaning, facts, dates, names, companies, skills, or any factual content
4. NEVER add: fake information, new skills, new projects, new experience, new certifications
5. NEVER remove: any user content, sections, or information
6. ONLY include corrections in the "corrections" list for lines that actually need improvement (e.g. lines with grammar errors, typos, weak verbs, or poor readability). Do NOT include lines that require no changes. If a line requires no correction, completely omit it from the returned corrections list.
7. Return ONLY the corrected JSON, nothing else

You will receive numbered lines of resume text. Format your response as a JSON object with a single key "corrections", containing a list of objects with "line_num" and "corrected" fields for ONLY the lines that were improved. Any omitted line indices will remain unchanged."""


class CorrectionService:
    def __init__(self, db: Session):
        self.db = db
        self.corrected_dir = "uploads/corrected"
        os.makedirs(self.corrected_dir, exist_ok=True)
        self.llm_service = LLMService()

    async def correct_resume(self, resume_id: int, user_id: int) -> dict:
        resume = self.db.query(Resume).filter(
            Resume.id == resume_id, Resume.user_id == user_id
        ).first()
        if not resume:
            raise NotFoundException("Resume not found")

        if not os.path.exists(resume.file_path):
            raise NotFoundException("Resume file not found on server")

        file_ext = resume.file_type.lower()
        if file_ext not in ("pdf", "docx"):
            raise BadRequestException("Only PDF and DOCX files can be corrected")

        if not self.llm_service.client:
            raise BadRequestException("OpenRouter API key is not configured. Set OPENROUTER_API_KEY in .env")

        try:
            if file_ext == "pdf":
                result = await self._correct_pdf(resume)
            else:
                result = await self._correct_docx(resume)

            # Automatically analyze the corrected resume and store it
            filename = result["corrected_pdf"] or result["corrected_docx"]
            if filename:
                output_path = os.path.join(self.corrected_dir, filename)
                
                from app.models.resume import Resume as ResumeModel
                from app.services.resume_service import ResumeService
                from app.services.analysis_service import AnalysisService
                from app.schemas.analysis import AnalysisCreate
                from app.models.analysis import Analysis as AnalysisModel

                # Parse the corrected resume text
                resume_service = ResumeService(self.db)
                ext = f".{resume.file_type}"
                extracted_data = resume_service._extract_data(output_path, ext)

                corrected_resume = ResumeModel(
                    user_id=user_id,
                    filename=filename,
                    original_filename=f"corrected_{resume.original_filename}",
                    file_path=output_path,
                    file_type=resume.file_type,
                    file_size=os.path.getsize(output_path),
                    extracted_data=extracted_data,
                )
                self.db.add(corrected_resume)
                self.db.commit()
                self.db.refresh(corrected_resume)

                # Fetch latest analysis job description to preserve context
                latest_analysis = self.db.query(AnalysisModel).filter(
                    AnalysisModel.resume_id == resume.id
                ).order_by(AnalysisModel.created_at.desc()).first()
                job_desc = latest_analysis.job_description if latest_analysis else None

                analysis_service = AnalysisService(self.db)
                analysis_data = AnalysisCreate(
                    resume_id=corrected_resume.id,
                    job_description=job_desc
                )
                corrected_analysis = await analysis_service.create_analysis(user_id, analysis_data)
                result["corrected_analysis"] = corrected_analysis

            return result
        except (BadRequestException, LLMAPIException):
            raise
        except Exception as e:
            logger.error(f"Correction failed for resume {resume_id}: {e}")
            raise BadRequestException(f"Correction failed: {str(e)}")

    async def _correct_pdf(self, resume: Resume) -> dict:
        file_path = resume.file_path
        doc = pymupdf.open(file_path)

        all_lines = []
        page_line_map = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            text_dict = page.get_text("dict", flags=pymupdf.TEXT_PRESERVE_WHITESPACE)

            lines_on_page = self._extract_pdf_lines(text_dict, page_num)
            for line_info in lines_on_page:
                idx = len(all_lines)
                page_line_map.append(line_info)
                all_lines.append({
                    "num": idx,
                    "text": line_info["text"],
                    "page": page_num,
                })

        if not all_lines:
            doc.close()
            raise BadRequestException("No text found in the PDF to correct")

        corrected_lines = await self._correct_text_batch(
            resume.id,
            [l["text"] for l in all_lines],
            resume.original_filename,
        )

        changes = []
        full_diff = []
        for i, (original, corrected) in enumerate(zip(all_lines, corrected_lines)):
            is_changed = original["text"].strip() != corrected.strip()
            if is_changed:
                changes.append({
                    "original": original["text"],
                    "corrected": corrected,
                    "line_num": i,
                })
            full_diff.append({
                "original": original["text"],
                "corrected": corrected,
                "changed": is_changed,
            })

        output_name = f"corrected_{uuid.uuid4().hex}.pdf"
        highlighted_name = output_name.replace("corrected_", "highlighted_")
        
        output_path = os.path.join(self.corrected_dir, output_name)
        highlighted_path = os.path.join(self.corrected_dir, highlighted_name)

        # Apply corrections without highlights first and save the clean file
        self._apply_pdf_corrections(doc, page_line_map, corrected_lines)
        doc.save(output_path)
        logger.info(f"PDF GENERATION - Resume ID: {resume.id}, Saved clean corrected PDF to: {output_path}")

        # Verify generated PDF text contains the corrections (Phase 10)
        try:
            doc_verify = pymupdf.open(output_path)
            extracted_verify = []
            for page in doc_verify:
                text_dict = page.get_text("dict", flags=pymupdf.TEXT_PRESERVE_WHITESPACE)
                lines_on_page = self._extract_pdf_lines(text_dict, page.number)
                extracted_verify.extend([l["text"] for l in lines_on_page])
            doc_verify.close()
            
            logger.info(f"PDF VERIFICATION - Resume ID: {resume.id}, Extracted {len(extracted_verify)} lines from generated PDF.")
            diffs_found = 0
            for idx, line_info in enumerate(all_lines):
                orig_text = line_info["text"]
                corr_text = corrected_lines[idx]
                if orig_text.strip() != corr_text.strip():
                    matched_verify = False
                    for line_v in extracted_verify:
                        if corr_text.strip() in line_v.strip() or line_v.strip() in corr_text.strip():
                            matched_verify = True
                            break
                    if not matched_verify:
                        logger.warning(f"PDF VERIFICATION WARNING - Line {idx} corrected to '{corr_text}' but not found in extracted PDF text!")
                        diffs_found += 1
            if diffs_found == 0:
                logger.info(f"PDF VERIFICATION SUCCESS - All corrected lines verified in the generated PDF.")
        except Exception as ve:
            logger.error(f"PDF VERIFICATION ERROR - Failed to verify generated PDF text: {ve}")
        
        # Apply highlights on top of the corrected document and save the highlighted file
        self._apply_pdf_highlights(doc, page_line_map, corrected_lines)
        doc.save(highlighted_path)
        doc.close()
        logger.info(f"PDF GENERATION - Resume ID: {resume.id}, Saved highlighted PDF to: {highlighted_path}")

        return {
            "corrected_pdf": output_name,
            "corrected_docx": None,
            "changes": changes,
            "full_diff": full_diff,
            "total_lines": len(all_lines),
            "changed_lines": len(changes),
        }

    async def _correct_docx(self, resume: Resume) -> dict:
        file_path = resume.file_path
        doc = Document(file_path)

        paragraphs_data = []
        for para_idx, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
            runs_text = []
            for run in para.runs:
                runs_text.append(run.text)
            full_text = "".join(runs_text)
            if full_text.strip():
                paragraphs_data.append({
                    "para_idx": para_idx,
                    "text": full_text,
                    "runs": [
                        {
                            "text": r.text,
                            "bold": r.bold,
                            "italic": r.italic,
                            "underline": r.underline,
                            "font_name": r.font.name if r.font else None,
                            "font_size": r.font.size.pt if r.font and r.font.size else None,
                            "font_color": str(r.font.color.rgb) if r.font and r.font.color and r.font.color.rgb else None,
                        }
                        for r in para.runs
                    ],
                })

        if not paragraphs_data:
            raise BadRequestException("No text found in the DOCX to correct")

        all_texts = [p["text"] for p in paragraphs_data]
        corrected_texts = await self._correct_text_batch(
            resume.id,
            all_texts,
            resume.original_filename,
        )

        changes = []
        full_diff = []
        for original, corrected in zip(paragraphs_data, corrected_texts):
            is_changed = original["text"].strip() != corrected.strip()
            if is_changed:
                changes.append({
                    "original": original["text"],
                    "corrected": corrected,
                    "para_idx": original["para_idx"],
                })
            full_diff.append({
                "original": original["text"],
                "corrected": corrected,
                "changed": is_changed,
            })

        output_name = f"corrected_{uuid.uuid4().hex}.docx"
        highlighted_name = output_name.replace("corrected_", "highlighted_")
        
        output_path = os.path.join(self.corrected_dir, output_name)
        highlighted_path = os.path.join(self.corrected_dir, highlighted_name)

        # Apply corrections without highlights first and save the clean file
        self._apply_docx_corrections(doc, paragraphs_data, corrected_texts)
        doc.save(output_path)
        
        # Apply highlights on top of the corrected document and save the highlighted file
        self._apply_docx_highlights(doc, paragraphs_data, corrected_texts)
        doc.save(highlighted_path)

        return {
            "corrected_pdf": None,
            "corrected_docx": output_name,
            "changes": changes,
            "full_diff": full_diff,
            "total_lines": len(paragraphs_data),
            "changed_lines": len(changes),
        }

    def _extract_pdf_lines(self, text_dict: dict, page_num: int) -> list:
        lines = []
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue

            block_lines = block.get("lines", [])
            for line in block_lines:
                spans = line.get("spans", [])
                if not spans:
                    continue

                line_text = "".join(s["text"] for s in spans).strip()
                if not line_text:
                    continue

                first_span = spans[0]
                bbox = line["bbox"]

                font_name = first_span.get("font", "helv")
                font_size = first_span.get("size", 10)
                color_int = first_span.get("color", 0)
                flags = first_span.get("flags", 0)

                lines.append({
                    "text": line_text,
                    "page": page_num,
                    "bbox": bbox,
                    "font_name": font_name,
                    "font_size": font_size,
                    "color": color_int,
                    "flags": flags,
                    "spans": spans,
                    "is_bold": bool(flags & 2 ** 4),
                    "is_italic": bool(flags & 2 ** 1),
                })

        lines.sort(key=lambda l: (l["bbox"][1], l["bbox"][0]))
        return lines

    def clean_and_parse_json(self, raw_text: str) -> dict:
        if not raw_text:
            raise ValueError("LLM response content was empty")
            
        text = raw_text.strip()
        
        # Strip markdown code fences if present at the start/end
        if text.startswith("```"):
            lines = text.split("\n")
            if len(lines) > 1:
                content_lines = []
                for line in lines[1:]:
                    if line.strip() == "```":
                        break
                    content_lines.append(line)
                text = "\n".join(content_lines).strip()
                
        # Also clean inline markdown wrapping/explanatory text if present
        first_brace = text.find("{")
        last_brace = text.rfind("}")
        
        if first_brace == -1 or last_brace == -1 or last_brace <= first_brace:
            raise ValueError("Could not find valid JSON object boundaries (curly braces)")
            
        json_candidate = text[first_brace:last_brace + 1]
        
        try:
            return json.loads(json_candidate)
        except json.JSONDecodeError as e:
            logger.warning(f"Direct JSON parsing failed ({e}). Attempting truncation repair...")
            # Try appending ']}' to close the truncated array and outer object
            repaired_json = json_candidate + "]}"
            try:
                return json.loads(repaired_json)
            except json.JSONDecodeError:
                # If direct append failed, search backward for the last valid item closing brace '}'
                # and attempt to slice and append ']}'
                curr_idx = last_brace
                while curr_idx > first_brace:
                    prev_brace = json_candidate.rfind("}", 0, curr_idx)
                    if prev_brace == -1:
                        break
                    repaired_candidate = json_candidate[:prev_brace + 1] + "]}"
                    try:
                        return json.loads(repaired_candidate)
                    except json.JSONDecodeError:
                        curr_idx = prev_brace
                
                # If all repairs failed, raise the original decoding error
                raise ValueError(f"Failed to parse or repair JSON: {e}")

    async def _correct_text_batch(self, resume_id: int, texts: list, filename: str) -> list:
        if not self.llm_service.client:
            return texts

        batch_size = 100
        all_corrected = []

        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]
            numbered_lines = []
            for j, text in enumerate(batch):
                numbered_lines.append(f"{j}: {text}")

            prompt = f"""Correct the following resume text lines from "{filename}".
Fix grammar, spelling, punctuation, improve action verbs and professional tone.
DO NOT change any factual content, names, dates, companies, or skills.
DO NOT add or remove any information.
Return a JSON object with a single key "corrections", containing a list of objects like {{"line_num": 0, "corrected": "corrected text"}}.
Only return the JSON object, nothing else.

Lines to correct:
{chr(10).join(numbered_lines)}"""

            import hashlib
            prompt_len = len(prompt)
            prompt_hash = hashlib.sha256(prompt.encode('utf-8')).hexdigest()
            first_1000 = prompt[:1000]
            logger.info(f"LLM REQUEST (Correction) - Resume ID: {resume_id}, Prompt Length: {prompt_len}, Prompt Hash: {prompt_hash}")
            logger.info(f"Prompt First 1000 characters:\n{first_1000}")

            import time
            max_retries = 3
            api_success = False
            parsed_data = None

            for attempt in range(max_retries):
                try:
                    result_text = await self.llm_service.generate_chat_completion(
                        messages=[
                            {"role": "system", "content": AI_SYSTEM_PROMPT},
                            {"role": "user", "content": prompt},
                        ],
                        response_format={"type": "json_object"},
                        temperature=0.3,
                        max_tokens=4000,
                    )
                    result_text = result_text.strip()
                    
                    # Log the raw response before parsing
                    logger.info(f"RAW LLM RESPONSE (Attempt {attempt + 1}/{max_retries}):\n{result_text}")
                    
                    if not result_text:
                        logger.warning("LLM response content was empty. Retrying...")
                        if attempt < max_retries - 1:
                            time.sleep(1)
                        continue

                    # Validate and parse JSON inside the retry loop
                    try:
                        parsed_data = self.clean_and_parse_json(result_text)
                        api_success = True
                        break
                    except ValueError as parse_err:
                        logger.warning(f"Invalid JSON returned from model on attempt {attempt + 1}: {parse_err}")
                        if attempt < max_retries - 1:
                            time.sleep(2)
                            continue
                        else:
                            raise BadRequestException(f"AI correction failed: Invalid JSON response from model after {max_retries} attempts. Raw details: {parse_err}")

                except Exception as e:
                    if isinstance(e, (BadRequestException, LLMAPIException)):
                        raise
                    raise handle_llm_exception(e)

            if not api_success or parsed_data is None:
                raise BadRequestException("AI correction failed: API did not respond successfully with valid JSON.")

            corrections = parsed_data.get("corrections", [])
            if not isinstance(corrections, list):
                corrections = []

            corrected_batch = list(batch)
            for item in corrections:
                if isinstance(item, dict):
                    idx = item.get("line_num", -1)
                    corrected_text = item.get("corrected", "")
                    if 0 <= idx < len(corrected_batch) and corrected_text:
                        corrected_batch[idx] = corrected_text

            all_corrected.extend(corrected_batch)

        while len(all_corrected) < len(texts):
            all_corrected.append(texts[len(all_corrected)])

        return all_corrected[:len(texts)]


    def _apply_pdf_corrections(self, doc, page_line_map, corrected_lines):
        pages_to_process = {}
        for i, line_info in enumerate(page_line_map):
            page_num = line_info["page"]
            if page_num not in pages_to_process:
                pages_to_process[page_num] = []
            pages_to_process[page_num].append((i, line_info))

        for page_num, line_indices in pages_to_process.items():
            page = doc[page_num]

            for idx, line_info in line_indices:
                corrected_text = corrected_lines[idx]
                original_text = line_info["text"]

                if original_text.strip() == corrected_text.strip():
                    continue

                bbox = line_info["bbox"]
                rect = pymupdf.Rect(bbox)

                mapped_font = self._map_font(line_info["font_name"])
                font_size = line_info["font_size"]
                color_int = line_info["color"]

                r = ((color_int >> 16) & 0xFF) / 255.0
                g = ((color_int >> 8) & 0xFF) / 255.0
                b = (color_int & 0xFF) / 255.0

                page.add_redact_annot(
                    rect,
                    text=None,
                    fill=(1, 1, 1),
                )

            page.apply_redactions()

            for idx, line_info in line_indices:
                corrected_text = corrected_lines[idx]
                original_text = line_info["text"]

                if original_text.strip() == corrected_text.strip():
                    continue

                bbox = line_info["bbox"]
                mapped_font = self._map_font(line_info["font_name"])
                font_size = line_info["font_size"]
                color_int = line_info["color"]

                r = ((color_int >> 16) & 0xFF) / 255.0
                g = ((color_int >> 8) & 0xFF) / 255.0
                b = (color_int & 0xFF) / 255.0

                origin_x = bbox[0]
                origin_y = bbox[3] - 2

                try:
                    page.insert_text(
                        pymupdf.Point(origin_x, origin_y),
                        corrected_text,
                        fontname=mapped_font,
                        fontsize=font_size,
                        color=(r, g, b),
                    )
                except Exception as e:
                    logger.warning(f"Failed to insert text at page {line_info['page']}: {e}")
                    try:
                        page.insert_text(
                            pymupdf.Point(origin_x, origin_y),
                            corrected_text,
                            fontname="helv",
                            fontsize=font_size,
                            color=(r, g, b),
                        )
                    except Exception as e2:
                        logger.error(f"Fallback text insert also failed: {e2}")

    def _apply_pdf_highlights(self, doc, page_line_map, corrected_lines):
        pages_to_process = {}
        for i, line_info in enumerate(page_line_map):
            page_num = line_info["page"]
            if page_num not in pages_to_process:
                pages_to_process[page_num] = []
            pages_to_process[page_num].append((i, line_info))

        for page_num, line_indices in pages_to_process.items():
            page = doc[page_num]
            for idx, line_info in line_indices:
                corrected_text = corrected_lines[idx]
                original_text = line_info["text"]

                if original_text.strip() == corrected_text.strip():
                    continue

                bbox = line_info["bbox"]
                try:
                    highlight = page.add_highlight_annot(pymupdf.Rect(bbox))
                    highlight.set_colors(stroke=(1.0, 0.9, 0.6))  # premium soft yellow color
                    highlight.update()
                except Exception as he:
                    logger.warning(f"Failed to add highlight annotation for line: {he}")

    def _apply_docx_corrections(self, doc, paragraphs_data, corrected_texts):
        para_idx = 0
        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            if para_idx >= len(paragraphs_data):
                break

            corrected_text = corrected_texts[para_idx]
            original_data = paragraphs_data[para_idx]
            para_idx += 1

            if original_data["text"].strip() == corrected_text.strip():
                continue

            runs = para.runs
            if not runs:
                para.text = corrected_text
                continue

            if len(runs) == 1:
                runs[0].text = corrected_text
                continue

            self._distribute_text_to_runs(runs, corrected_text, original_data["runs"])

    def _apply_docx_highlights(self, doc, paragraphs_data, corrected_texts):
        from docx.enum.text import WD_COLOR_INDEX
        para_idx = 0
        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            if para_idx >= len(paragraphs_data):
                break

            corrected_text = corrected_texts[para_idx]
            original_data = paragraphs_data[para_idx]
            para_idx += 1

            if original_data["text"].strip() == corrected_text.strip():
                continue

            for run in para.runs:
                try:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                except Exception:
                    pass

    def _distribute_text_to_runs(self, runs, corrected_text, runs_meta):
        if not runs or not corrected_text:
            return

        total_original_len = sum(len(r.text) for r in runs)
        if total_original_len == 0:
            runs[0].text = corrected_text
            for r in runs[1:]:
                r.text = ""
            return

        remaining = corrected_text
        for i, run in enumerate(runs):
            if i == len(runs) - 1:
                run.text = remaining
            else:
                original_len = len(run.text)
                ratio = original_len / total_original_len
                chars_to_take = max(1, int(len(corrected_text) * ratio))

                split_point = self._find_best_split_point(remaining, chars_to_take)
                run.text = remaining[:split_point]
                remaining = remaining[split_point:]

    def _find_best_split_point(self, text, target_pos):
        if target_pos >= len(text):
            return len(text)

        best = target_pos
        for offset in range(min(5, target_pos)):
            for pos in [target_pos + offset, target_pos - offset]:
                if 0 <= pos <= len(text):
                    if pos < len(text) and text[pos] == " ":
                        return pos + 1
                    if pos > 0 and text[pos - 1] == " ":
                        return pos
                    best = pos

        return best

    def _map_font(self, font_name: str) -> str:
        if not font_name:
            return "helv"

        normalized = font_name.lower().strip()

        for key, mapped in FONT_MAP.items():
            if key in normalized:
                return mapped

        if "bold" in normalized and "italic" in normalized:
            return "hebi"
        if "bold" in normalized:
            return "hebo"
        if "italic" in normalized or "oblique" in normalized:
            return "heit"

        return "helv"

    def download_corrected_file(self, filename: str) -> str:
        clean_filename = os.path.basename(filename)
        file_path = os.path.join(self.corrected_dir, clean_filename)
        if not os.path.exists(file_path):
            raise NotFoundException("Corrected file not found")
        return file_path

    def cleanup_old_files(self, max_age_hours: int = 24):
        import time
        now = time.time()
        cutoff = now - (max_age_hours * 3600)

        for filename in os.listdir(self.corrected_dir):
            filepath = os.path.join(self.corrected_dir, filename)
            if os.path.isfile(filepath) and os.path.getmtime(filepath) < cutoff:
                try:
                    os.remove(filepath)
                except Exception:
                    pass
